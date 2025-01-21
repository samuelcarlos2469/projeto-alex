import os
from flask import Flask, render_template, request, redirect, url_for, send_from_directory
import cv2
from PIL import Image
import google.generativeai as genai
from docx import Document
import PyPDF2

genai.configure(api_key="KEY_API")
model = genai.GenerativeModel("gemini-1.5-flash")

app = Flask(__name__, static_folder='static')

UPLOAD_FOLDER = 'uploads'
RESULT_FOLDER = 'uploads'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(RESULT_FOLDER):
    os.makedirs(RESULT_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULT_FOLDER'] = RESULT_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/img_to_doc')
def img_to_doc():
    return render_template('img_to_doc.html')

@app.route('/pdf_to_doc')
def pdf_to_doc():
    return render_template('pdf_to_doc.html')

@app.route('/upload_image', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return "Nenhum arquivo enviado", 400

    file = request.files['file']
    if file.filename == '':
        return "Nenhum arquivo selecionado", 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        image = cv2.imread(filepath)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 102, 255, cv2.THRESH_BINARY_INV)
        processed_image_path = os.path.join(app.config['UPLOAD_FOLDER'], "imagem_processada.png")
        cv2.imwrite(processed_image_path, thresh)

        processed_image = Image.open(processed_image_path)
        response = model.generate_content([
            "O que está escrito? Apenas transcreva o texto sem comentar nada",
            processed_image
        ])

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        text_doc = response.text
        doc.add_paragraph(text_doc)

        doc_path = os.path.join(app.config['RESULT_FOLDER'], "document.docx")
        doc.save(doc_path)

        return redirect(url_for('download_file', filename="document.docx"))

    except Exception as e:
        return f"Erro ao processar a imagem: {e}", 500

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return "Nenhum arquivo enviado", 400

    file = request.files['file']
    if file.filename == '':
        return "Nenhum arquivo selecionado", 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        with open(filepath, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            doc = Document()

            for page in reader.pages:
                text = page.extract_text()
                doc.add_paragraph(text)

            doc_path = os.path.join(app.config['RESULT_FOLDER'], "converted_document.docx")
            doc.save(doc_path)

        return redirect(url_for('download_file', filename="converted_document.docx"))

    except Exception as e:
        return f"Erro ao processar o PDF: {e}", 500

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['RESULT_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
